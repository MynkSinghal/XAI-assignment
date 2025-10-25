#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate the IEEE-style Explainable AI report for the heart disease project.

This utility compiles the analytical assets produced by the data pipeline into a
publication-ready Microsoft Word document. The script expects the latest model
artifacts to be present in the ``visuals/`` directory (figures, CSV metrics,
SHAP export) and will recreate ``XAI_Report_Mayank_Utkarsh.docx`` under the
``report/`` directory.

Usage
-----
Run the pipeline first so that metrics, fairness diagnostics, and SHAP plots are
up to date. Then regenerate the report with:

    python report/generate_report.py

Optional arguments allow overriding the output path:

    python report/generate_report.py --output path/to/custom_report.docx

The generated document follows IEEE margin and typography conventions and is
intended for light manual editing (e.g., minor copyedits or pagination tweaks)
within Microsoft Word after creation.
"""

from __future__ import annotations

import argparse
from datetime import datetime
from pathlib import Path
from typing import List, Sequence

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

IEEE_FONT = "Times New Roman"
BODY_FONT_SIZE = Pt(10)
HEADING_FONT_SIZE = Pt(12)
TITLE_FONT_SIZE = Pt(16)
MARGIN_INCHES = 0.75
REPORT_FILENAME = "XAI_Report_Mayank_Utkarsh.docx"

REPORT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = REPORT_DIR.parent
VISUALS_DIR = PROJECT_ROOT / "visuals"


def ensure_assets_exist(paths: Sequence[Path]) -> None:
    """Validate that all required files are present before composing the report."""
    missing = [path for path in paths if not path.exists()]
    if missing:
        missing_str = "\n".join(str(path) for path in missing)
        raise FileNotFoundError(
            "The following required assets are missing. Run the data pipeline "
            "before regenerating the report:\n" + missing_str
        )


def configure_document(document: Document) -> None:
    """Apply IEEE-aligned defaults for margins, fonts, and paragraph spacing."""
    for section in document.sections:
        section.top_margin = Inches(MARGIN_INCHES)
        section.bottom_margin = Inches(MARGIN_INCHES)
        section.left_margin = Inches(MARGIN_INCHES)
        section.right_margin = Inches(MARGIN_INCHES)

    normal_style = document.styles["Normal"]
    normal_style.font.name = IEEE_FONT
    normal_style.font.size = BODY_FONT_SIZE
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.line_spacing = 1.0

    heading_style = document.styles["Heading 1"]
    heading_style.font.name = IEEE_FONT
    heading_style.font.size = HEADING_FONT_SIZE
    heading_style.font.bold = True
    heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_style.paragraph_format.space_before = Pt(12)
    heading_style.paragraph_format.space_after = Pt(6)

    title_style = document.styles["Title"]
    title_style.font.name = IEEE_FONT
    title_style.font.size = TITLE_FONT_SIZE
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_title_page(document: Document, authors: List[str]) -> None:
    """Insert a centered title block with authorship metadata."""
    title_para = document.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("Explainable AI for Heart Disease Prediction")
    title_run.font.name = IEEE_FONT
    title_run.font.size = TITLE_FONT_SIZE
    title_run.bold = True

    subtitle_para = document.add_paragraph("IEEE-Style Technical Report")
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle_para.runs:
        run.font.name = IEEE_FONT
        run.font.size = Pt(12)
        run.italic = True

    authors_para = document.add_paragraph(", ".join(authors))
    authors_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in authors_para.runs:
        run.font.name = IEEE_FONT
        run.font.size = Pt(11)

    group_para = document.add_paragraph("Group 07 · XAI Heart Disease Initiative")
    group_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in group_para.runs:
        run.font.name = IEEE_FONT
        run.font.size = Pt(10)

    date_para = document.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in date_para.runs:
        run.font.name = IEEE_FONT
        run.font.size = Pt(10)

    document.add_page_break()


def add_section_heading(document: Document, text: str) -> None:
    """Create an IEEE-aligned section heading."""
    heading = document.add_heading(text, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.name = IEEE_FONT
        run.font.size = HEADING_FONT_SIZE
        run.bold = True


def add_justified_paragraph(document: Document, text: str) -> None:
    """Append a justified paragraph with IEEE body styling."""
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in paragraph.runs:
        run.font.name = IEEE_FONT
        run.font.size = BODY_FONT_SIZE


def add_table(document: Document, headers: Sequence[str], rows: List[Sequence[str]], caption: str, table_number: int) -> int:
    """Insert a formatted table and caption, returning the next table index."""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        for run in header_cells[idx].paragraphs[0].runs:
            run.font.name = IEEE_FONT
            run.font.size = BODY_FONT_SIZE
            run.bold = True
        header_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_values in rows:
        row_cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            row_cells[idx].text = str(value)
            paragraph = row_cells[idx].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.name = IEEE_FONT
                run.font.size = BODY_FONT_SIZE

    caption_para = document.add_paragraph(f"Table {table_number}. {caption}")
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption_para.runs:
        run.font.name = IEEE_FONT
        run.font.size = BODY_FONT_SIZE
        run.italic = True

    return table_number + 1


def add_figure(document: Document, image_path: Path, caption: str, figure_number: int, width_inches: float = 5.5) -> int:
    """Insert a centered figure with caption, returning the next figure index."""
    document.add_picture(str(image_path), width=Inches(width_inches))
    picture_para = document.paragraphs[-1]
    picture_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption_para = document.add_paragraph(f"Figure {figure_number}. {caption}")
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption_para.runs:
        run.font.name = IEEE_FONT
        run.font.size = BODY_FONT_SIZE
        run.italic = True

    return figure_number + 1


def percent(value: float) -> str:
    """Format a decimal proportion as a percentage string with one decimal place."""
    return f"{value * 100:.1f}%"


def build_fairness_paragraphs(
    metrics_df: pd.DataFrame,
    fairness_by_sex: pd.DataFrame,
    fairness_by_age: pd.DataFrame,
    top_features: List[str],
) -> List[str]:
    """Compose fairness discussion paragraphs and enforce the 500–700 word policy."""
    rf_metrics = metrics_df.loc[metrics_df["Model"] == "Random Forest"].iloc[0]
    lr_metrics = metrics_df.loc[metrics_df["Model"] == "Logistic Regression"].iloc[0]

    sex_lookup = {
        int(row.Sex): row for row in fairness_by_sex.itertuples(index=False)
    }
    if 0 not in sex_lookup or 1 not in sex_lookup:
        raise ValueError("Fairness by sex metrics must contain both sex categories (0 and 1).")

    female = sex_lookup[0]
    male = sex_lookup[1]
    precision_gap = abs(female.Precision - male.Precision) * 100

    age_lookup = {
        str(row.Age_Group): row for row in fairness_by_age.itertuples(index=False)
    }
    try:
        younger = age_lookup["<50"]
        middle = age_lookup["50-60"]
        senior = age_lookup["60+"]
    except KeyError as exc:  # pragma: no cover - defensive guard
        raise ValueError("Fairness by age metrics must include '<50', '50-60', and '60+' groups.") from exc

    paragraphs = [
        (
            "Ensuring equitable deployment of the heart disease classifier is essential for integrating the "
            "model into clinical triage pathways without amplifying disparities. The Random Forest configuration "
            f"selected for deployment combines an accuracy of {rf_metrics['Accuracy']:.3f} with a ROC-AUC of "
            f"{rf_metrics['ROC_AUC']:.3f}, outperforming the logistic regression baseline while preserving sensitivity "
            "to high-risk patients. Fairness assessment therefore focuses on verifying that this uplift does not "
            "concentrate benefits into a narrow group. Table 2 consolidates the demographic diagnostics that ground "
            "this review. We draw on segment-level confusion statistics, the SHAP attributions that expose feature "
            "leverage [1], and cost-oriented considerations from responsible AI literature [4] to guide remediation planning."
        ),
        (
            "Sex-specific evaluation indicates that the model is slightly more conservative for male patients than for "
            "female patients. Among female records (n="
            f"{int(female.Count)}), the classifier achieved {percent(female.Accuracy)} accuracy, {percent(female.Precision)} precision, "
            f"{percent(female.Recall)} recall, and an F1-score of {percent(female.F1)}, indicating near-perfect specificity but a handful of missed positives. "
            f"Male records (n={int(male.Count)}) show {percent(male.Accuracy)} accuracy, {percent(male.Precision)} precision, "
            f"{percent(male.Recall)} recall, and {percent(male.F1)} F1, illustrating the inverse pattern of excellent sensitivity with additional false positives. "
            f"The {precision_gap:.1f}% gap in precision suggests that threshold calibration or re-weighting may be required to harmonize downstream burdens between sexes, "
            "even though neither group experiences catastrophic performance loss."
        ),
        (
            "Age-stratified diagnostics reveal a second axis of uneven performance that must be managed. Younger-than-fifty patients "
            f"(n={int(younger.Count)}) achieved {percent(younger.Accuracy)} accuracy with a perfect F1-score of {percent(younger.F1)}, while the 50–60 cohort "
            f"(n={int(middle.Count)}) retained {percent(middle.Accuracy)} accuracy and {percent(middle.F1)} F1. Performance declines for seniors aged sixty and above "
            f"(n={int(senior.Count)}), whose outcomes drop to {percent(senior.Accuracy)} accuracy and {percent(senior.F1)} F1, implying reduced recall in this subgroup. "
            "Potential drivers include class imbalance and the dominant influence of exercise-induced metrics, which may capture younger physiology more faithfully. "
            "These findings motivate age-aware post-hoc calibration, additional data collection, or targeted clinician review queues when the model evaluates older patients."
        ),
        (
            "Interpretability analysis compares the transparency profiles of the two candidate models. Logistic regression offers "
            "inherently linear coefficients that clinicians can inspect, yet its {0:.1f}% accuracy and {1:.1f}% F1-score underscore limited capacity to model "
            "feature interactions. The Random Forest, by contrast, leverages ensemble diversity while remaining explainable through SHAP value decomposition [1]. "
            "The summary visualization shows {2}, {3}, and {4} exerting the largest marginal contributions, with {5} and {6} providing additional nuance around vascular stress. "
            "Dependence trends illustrate how moderate elevations in these attributes shift predictions, enabling practitioners to reconcile the tree-based model’s "
            "complexity with human reasoning.".format(
                lr_metrics["Accuracy"] * 100,
                lr_metrics["F1"] * 100,
                top_features[0],
                top_features[1],
                top_features[2],
                top_features[3],
                top_features[4],
            )
        ),
        (
            "These fairness observations surface tangible trade-offs among accuracy, sensitivity, and equitable treatment. Tightening the decision threshold to "
            "raise male precision would simultaneously reduce recall, potentially missing positive cases that the care team must triage promptly. Conversely, lowering "
            "the threshold to assist older patients could inflate false positives for younger cohorts and burden diagnostic services. Cost-sensitive retraining or "
            "group-specific calibration can balance these pressures, but teams must evaluate operational capacity and regulatory expectations before implementation. "
            "Figures 1–3 help clinicians visualize the implications by pairing distributional errors with probability contours along the ROC operating curve [2]."
        ),
        (
            "To operationalize these insights responsibly, we recommend instituting a quarterly fairness dashboard that refreshes Tables 1 and 2 with current hospital data, "
            "supported by automated alerts when segmental precision deviates by more than five percentage points. Complementary measures include augmenting training data for "
            "seniors, exploring monotonic gradient boosting as an interpretable alternative, and incorporating clinician feedback loops that flag misclassified cases for retraining. "
            f"We also advise documenting SHAP-based explanation templates so staff can communicate the influence of {top_features[0]} and {top_features[1]} on individual predictions. "
            "Together these steps sustain transparency, encourage equitable care, and meet governance expectations around algorithmic accountability."
        ),
    ]

    word_count = sum(len(paragraph.split()) for paragraph in paragraphs)
    if not 500 <= word_count <= 700:
        raise ValueError(
            f"Fairness discussion length is {word_count} words; update the narrative to stay within the 500–700 word range."
        )

    return paragraphs


def load_project_artifacts() -> dict:
    """Load metrics, fairness diagnostics, and feature attributions from disk."""
    metrics_df = pd.read_csv(VISUALS_DIR / "metrics.csv").dropna(how="all")
    fairness_by_sex = pd.read_csv(VISUALS_DIR / "fairness_by_sex.csv").dropna(how="all")
    fairness_by_age = pd.read_csv(VISUALS_DIR / "fairness_by_age.csv").dropna(how="all")
    importance_df = pd.read_csv(VISUALS_DIR / "feature_importance_data.csv").dropna(how="all")

    # Sort feature importance descending to extract top contributors.
    importance_df = importance_df.sort_values("Importance", ascending=False).reset_index(drop=True)

    return {
        "metrics": metrics_df,
        "fairness_sex": fairness_by_sex,
        "fairness_age": fairness_by_age,
        "importance": importance_df,
    }


def build_model_metrics_table(metrics_df: pd.DataFrame) -> List[Sequence[str]]:
    """Transform metric values into table-friendly rows."""
    ordered_columns = ["Model", "Accuracy", "Precision", "Recall", "F1", "ROC_AUC"]
    rows: List[Sequence[str]] = []
    for _, model_row in metrics_df[ordered_columns].iterrows():
        rows.append(
            [
                model_row["Model"],
                percent(model_row["Accuracy"]),
                percent(model_row["Precision"]),
                percent(model_row["Recall"]),
                percent(model_row["F1"]),
                f"{model_row['ROC_AUC']:.3f}",
            ]
        )
    return rows


def build_fairness_table_rows(fairness_by_sex: pd.DataFrame, fairness_by_age: pd.DataFrame) -> List[Sequence[str]]:
    """Combine sex- and age-based fairness metrics into a single table."""
    rows: List[Sequence[str]] = []

    for row in fairness_by_sex.itertuples(index=False):
        group_label = "Female" if int(row.Sex) == 0 else "Male"
        rows.append(
            [
                "Sex",
                group_label,
                str(int(row.Count)),
                percent(row.Accuracy),
                percent(row.Precision),
                percent(row.Recall),
                percent(row.F1),
            ]
        )

    for row in fairness_by_age.itertuples(index=False):
        rows.append(
            [
                "Age",
                str(row.Age_Group),
                str(int(row.Count)),
                percent(row.Accuracy),
                percent(row.Precision),
                percent(row.Recall),
                percent(row.F1),
            ]
        )

    return rows


def compose_report(output_path: Path) -> None:
    """Create the IEEE-formatted report leveraging pipeline outputs."""
    assets = load_project_artifacts()
    metrics_df: pd.DataFrame = assets["metrics"]
    fairness_by_sex: pd.DataFrame = assets["fairness_sex"]
    fairness_by_age: pd.DataFrame = assets["fairness_age"]
    importance_df: pd.DataFrame = assets["importance"]

    rf_metrics = metrics_df.loc[metrics_df["Model"] == "Random Forest"].iloc[0]
    lr_metrics = metrics_df.loc[metrics_df["Model"] == "Logistic Regression"].iloc[0]
    top_features = importance_df.head(5)["Feature"].tolist()

    fairness_paragraphs = build_fairness_paragraphs(metrics_df, fairness_by_sex, fairness_by_age, top_features)

    document = Document()
    configure_document(document)
    add_title_page(document, ["Mayank Singhal", "Utkarsh Shukla"])

    # Abstract
    add_section_heading(document, "Abstract")
    abstract_text = (
        "Explainable AI techniques were applied to the UCI Cleveland heart disease dataset [3] to evaluate logistic regression and Random Forest "
        f"classifiers. After standardized preprocessing and one-hot encoding, the Random Forest achieved {percent(rf_metrics['Accuracy'])} accuracy with a ROC-AUC of {rf_metrics['ROC_AUC']:.3f}, "
        f"surpassing the logistic regression baseline at {percent(lr_metrics['Accuracy'])} accuracy. The pipeline generated calibrated probability curves, confusion matrices, and SHAP-based explanations [1] to ensure "
        "clinical interpretability. Fairness diagnostics across sex and age groups, summarized in Table 2, highlight precision and recall deltas that inform governance actions. "
        "The resulting reporting workflow satisfies IEEE formatting expectations and can be regenerated from curated assets for review, audit, and targeted iteration."
    )
    add_justified_paragraph(document, abstract_text)

    # Methodology
    add_section_heading(document, "Methodology")
    methodology_paragraph_1 = (
        "Data ingestion begins by retrieving the Cleveland subset of the UCI heart disease dataset [3] and applying cleaning operations to impute missing values, "
        "binarize the disease label, and engineer balanced train/test splits. Numerical and categorical features are organized via one-hot encoding, and the processed dataset "
        "is retained for auditability in downstream reviews."
    )
    methodology_paragraph_2 = (
        "Model training uses scikit-learn implementations of logistic regression and Random Forest classifiers [2], with hyperparameters aligned to the team’s pipeline defaults. "
        "Evaluation spans accuracy, precision, recall, F1-score, and ROC-AUC. Explainability artifacts rely on SHAP TreeExplainer outputs [1] to build summary perspectives, while fairness "
        "watchdog routines compute performance slices across age and sex cohorts consistent with governance recommendations [4]. The reporting script composes these outputs into publication-ready assets."
    )
    add_justified_paragraph(document, methodology_paragraph_1)
    add_justified_paragraph(document, methodology_paragraph_2)

    # Results and Visualizations
    add_section_heading(document, "Results and Visualizations")
    accuracy_delta = (rf_metrics["Accuracy"] - lr_metrics["Accuracy"]) * 100
    auc_delta = rf_metrics["ROC_AUC"] - lr_metrics["ROC_AUC"]
    results_paragraph_1 = (
        "Table 1 summarizes evaluation metrics for the two candidate models. The Random Forest maintains "
        f"{percent(rf_metrics['Accuracy'])} accuracy and {percent(rf_metrics['F1'])} F1, relative to logistic regression’s {percent(lr_metrics['Accuracy'])} accuracy and {percent(lr_metrics['F1'])} F1. "
        f"It also delivers the strongest ROC-AUC at {rf_metrics['ROC_AUC']:.3f}, exceeding the baseline by {auc_delta:.3f} and offering a {accuracy_delta:.1f}% gain in accuracy. These margins underpin the choice of the ensemble as the "
        "production-ready configuration."
    )
    results_paragraph_2 = (
        "Figure 1 visualizes the Random Forest confusion matrix, clarifying the balance between false positives and false negatives. Figure 2 presents ROC curves for both "
        "models across decision thresholds, while Figure 3 summarizes SHAP contributions, confirming that features such as "
        f"{top_features[0]} and {top_features[1]} dominate risk scoring and align with established cardiovascular indicators."
    )
    add_justified_paragraph(document, results_paragraph_1)

    table_counter = 1
    table_headers = ["Model", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC"]
    table_rows = build_model_metrics_table(metrics_df)
    table_counter = add_table(
        document,
        table_headers,
        table_rows,
        "Model performance comparison on the Cleveland heart disease test split.",
        table_counter,
    )

    add_justified_paragraph(document, results_paragraph_2)

    figure_counter = 1
    figures = [
        (
            VISUALS_DIR / "confusion_matrix.png",
            "Confusion matrix for Random Forest predictions on the test cohort.",
        ),
        (
            VISUALS_DIR / "roc_curves.png",
            "Receiver operating characteristic curves for logistic regression and Random Forest classifiers.",
        ),
        (
            VISUALS_DIR / "shap_summary.png",
            "SHAP summary plot highlighting feature influence on Random Forest predictions.",
        ),
    ]

    for image_path, caption in figures:
        figure_counter = add_figure(document, image_path, caption, figure_counter)

    # Interpretability and Fairness Discussion
    add_section_heading(document, "Interpretability and Fairness Discussion")
    # First fairness paragraph mentions Table 2; add it, then continue with remaining paragraphs.
    add_justified_paragraph(document, fairness_paragraphs[0])

    fairness_headers = ["Segment", "Group", "Count", "Accuracy", "Precision", "Recall", "F1"]
    fairness_rows = build_fairness_table_rows(fairness_by_sex, fairness_by_age)
    table_counter = add_table(
        document,
        fairness_headers,
        fairness_rows,
        "Fairness metrics across sex and age cohorts for the Random Forest model.",
        table_counter,
    )

    for paragraph in fairness_paragraphs[1:]:
        add_justified_paragraph(document, paragraph)

    # References
    add_section_heading(document, "References")
    references = [
        "[1] S. M. Lundberg and S.-I. Lee, \"A Unified Approach to Interpreting Model Predictions,\" Advances in Neural Information Processing Systems, vol. 30, pp. 4765-4774, 2017.",
        "[2] F. Pedregosa et al., \"Scikit-learn: Machine Learning in Python,\" Journal of Machine Learning Research, vol. 12, pp. 2825-2830, 2011.",
        "[3] D. Dua and C. Graff, \"UCI Machine Learning Repository,\" University of California, Irvine, 2019. [Online]. Available: https://archive.ics.uci.edu/ml",
        "[4] M. Mitchell et al., \"Model Cards for Model Reporting,\" Proc. FAT*, pp. 220-229, 2019.",
    ]
    for reference in references:
        add_justified_paragraph(document, reference)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def parse_args() -> argparse.Namespace:
    """Parse command-line arguments for the report generator."""
    parser = argparse.ArgumentParser(description="Generate the IEEE-style heart disease XAI report.")
    parser.add_argument(
        "--output",
        type=Path,
        default=REPORT_DIR / REPORT_FILENAME,
        help="Optional path for the generated DOCX report.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    output_path: Path = args.output.resolve()

    required_assets = [
        VISUALS_DIR / "metrics.csv",
        VISUALS_DIR / "fairness_by_sex.csv",
        VISUALS_DIR / "fairness_by_age.csv",
        VISUALS_DIR / "feature_importance_data.csv",
        VISUALS_DIR / "confusion_matrix.png",
        VISUALS_DIR / "roc_curves.png",
        VISUALS_DIR / "shap_summary.png",
    ]
    ensure_assets_exist(required_assets)

    compose_report(output_path)
    print(f"IEEE report generated at {output_path}. You may now open the file in Word for light editing if desired.")


if __name__ == "__main__":
    main()
